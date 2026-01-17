from io import BytesIO
from docx import Document

def replace_placeholders_docx_bytes(template_path: str, mapping: dict[str, str]) -> bytes:
    """
    Carga plantilla DOCX, reemplaza placeholders y devuelve el DOCX final en bytes.
    """
    doc = Document(template_path)

    def _replace_in_paragraph(p):
        # Reemplazo simple por runs (suficiente si placeholders no están partidos en runs)
        text = p.text
        for k, v in mapping.items():
            text = text.replace(k, v or "")
        # reescribe el párrafo
        if text != p.text:
            p.clear()
            p.add_run(text)

    # párrafos
    for p in doc.paragraphs:
        _replace_in_paragraph(p)

    # tablas
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()
